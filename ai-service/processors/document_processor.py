"""
Document Processor - Text extraction from PDF and DOCX
Supports: PDF, DOCX, TXT
"""

import os
from typing import Dict, Any, List, Optional
from pathlib import Path

from .base import BaseProcessor, ProcessingResult, ProcessorError, ProcessorProvider
from services.llm_service import get_llm_service
from config import settings


class DocumentProcessor(BaseProcessor):
    """
    Document processor for text extraction and analysis

    Current implementation:
    - PDF: PyPDF2 + pdfplumber (robust extraction)
    - DOCX: python-docx
    - TXT: Direct reading
    - Analysis: Gemini Flash (free tier)

    Future support:
    - Google Docs API integration
    - Advanced invoice/receipt parsing
    - Table extraction
    """

    SUPPORTED_FORMATS = ['pdf', 'docx', 'doc', 'txt', 'md']

    def __init__(self, provider: ProcessorProvider = ProcessorProvider.GEMINI_API):
        """Initialize document processor"""
        super().__init__(provider)
        self.llm = get_llm_service()

    def get_supported_formats(self) -> List[str]:
        """Get supported document formats"""
        return self.SUPPORTED_FORMATS

    def validate_file(self, file_path: str) -> bool:
        """Validate document file"""
        path = Path(file_path)
        if not path.exists():
            return False

        extension = path.suffix.lstrip('.').lower()
        return extension in self.SUPPORTED_FORMATS

    async def process(
        self,
        file_path: str,
        options: Optional[Dict[str, Any]] = None
    ) -> ProcessingResult:
        """
        Process document file

        Args:
            file_path: Path to document file
            options: Optional parameters:
                - extract_tables: Whether to extract tables
                - detect_invoice: Try to detect invoice/receipt fields
                - max_pages: Maximum pages to process (for large PDFs)

        Returns:
            ProcessingResult with extracted text and analysis
        """
        try:
            options = options or {}

            # Validate file
            if not self.validate_file(file_path):
                raise ProcessorError(f"Invalid or unsupported document: {file_path}")

            # Extract text based on file type
            path = Path(file_path)
            extension = path.suffix.lstrip('.').lower()

            if extension == 'pdf':
                extracted_text, metadata = await self._extract_from_pdf(file_path, options)
            elif extension in ['docx', 'doc']:
                extracted_text, metadata = await self._extract_from_docx(file_path, options)
            elif extension in ['txt', 'md']:
                extracted_text, metadata = await self._extract_from_text(file_path, options)
            else:
                raise ProcessorError(f"Unsupported document format: {extension}")

            # Detect document type
            doc_type = self._detect_document_type(extracted_text, metadata)

            # Analyze content with LLM
            analysis = await self.llm.analyze_content(extracted_text, f"{doc_type} document")

            # Extract invoice/receipt data if detected
            invoice_data = None
            if doc_type in ['invoice', 'receipt', 'bill']:
                invoice_data = await self.llm.extract_invoice_details(extracted_text)

            return ProcessingResult(
                success=True,
                raw_text=extracted_text,
                summary=analysis.get("summary"),
                key_points=self._extract_key_points(extracted_text, metadata),
                extracted_tasks=analysis.get("tasks", []),
                metadata={
                    "provider": self.provider.value,
                    "file_type": extension,
                    "document_type": doc_type,
                    "page_count": metadata.get("page_count"),
                    "word_count": len(extracted_text.split()),
                    "character_count": len(extracted_text),
                    "has_tables": metadata.get("has_tables", False),
                    "invoice_data": invoice_data,
                    "is_resource": analysis.get("is_resource", False),
                    "suggested_projects": analysis.get("projects", []),
                    "suggested_areas": analysis.get("areas", []),
                },
                confidence_score=analysis.get("confidence", 0.85),
            )

        except Exception as e:
            return ProcessingResult(
                success=False,
                error=f"Document processing failed: {str(e)}"
            )

    async def _extract_from_pdf(
        self,
        pdf_path: str,
        options: Dict[str, Any]
    ) -> tuple[str, Dict[str, Any]]:
        """
        Extract text from PDF using PyPDF2 and pdfplumber

        Args:
            pdf_path: Path to PDF file
            options: Extraction options

        Returns:
            Tuple of (extracted_text, metadata)
        """
        try:
            import pdfplumber
            from PyPDF2 import PdfReader

            text_parts = []
            metadata = {
                "page_count": 0,
                "has_tables": False,
            }

            max_pages = options.get('max_pages', 100)  # Limit for large PDFs

            # Use pdfplumber for better text extraction
            with pdfplumber.open(pdf_path) as pdf:
                metadata["page_count"] = len(pdf.pages)

                for i, page in enumerate(pdf.pages):
                    if i >= max_pages:
                        break

                    # Extract text
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)

                    # Check for tables
                    if options.get('extract_tables') and page.extract_tables():
                        metadata["has_tables"] = True

            extracted_text = "\n\n".join(text_parts)

            # Fallback to PyPDF2 if pdfplumber failed
            if not extracted_text.strip():
                reader = PdfReader(pdf_path)
                metadata["page_count"] = len(reader.pages)

                for i, page in enumerate(reader.pages):
                    if i >= max_pages:
                        break
                    text_parts.append(page.extract_text())

                extracted_text = "\n\n".join(text_parts)

            return extracted_text, metadata

        except ImportError:
            raise ProcessorError(
                "PDF libraries not installed. Install with: "
                "pip install pdfplumber PyPDF2"
            )
        except Exception as e:
            raise ProcessorError(f"PDF extraction failed: {str(e)}")

    async def _extract_from_docx(
        self,
        docx_path: str,
        options: Dict[str, Any]
    ) -> tuple[str, Dict[str, Any]]:
        """
        Extract text from DOCX file

        Args:
            docx_path: Path to DOCX file
            options: Extraction options

        Returns:
            Tuple of (extracted_text, metadata)
        """
        try:
            from docx import Document

            doc = Document(docx_path)

            # Extract paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

            # Extract tables if requested
            table_text = []
            if options.get('extract_tables') and doc.tables:
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join(cell.text for cell in row.cells)
                        table_text.append(row_text)

            metadata = {
                "page_count": None,  # DOCX doesn't have pages concept
                "paragraph_count": len(paragraphs),
                "has_tables": len(doc.tables) > 0,
            }

            extracted_text = "\n\n".join(paragraphs)
            if table_text:
                extracted_text += "\n\n[Tables]\n" + "\n".join(table_text)

            return extracted_text, metadata

        except ImportError:
            raise ProcessorError(
                "python-docx not installed. Install with: pip install python-docx"
            )
        except Exception as e:
            raise ProcessorError(f"DOCX extraction failed: {str(e)}")

    async def _extract_from_text(
        self,
        text_path: str,
        options: Dict[str, Any]
    ) -> tuple[str, Dict[str, Any]]:
        """
        Extract text from plain text file

        Args:
            text_path: Path to text file
            options: Extraction options

        Returns:
            Tuple of (extracted_text, metadata)
        """
        try:
            with open(text_path, 'r', encoding='utf-8') as f:
                text = f.read()

            metadata = {
                "page_count": None,
                "line_count": len(text.split('\n')),
                "has_tables": False,
            }

            return text, metadata

        except Exception as e:
            # Try different encodings
            try:
                with open(text_path, 'r', encoding='latin-1') as f:
                    text = f.read()

                metadata = {
                    "page_count": None,
                    "line_count": len(text.split('\n')),
                    "has_tables": False,
                }

                return text, metadata

            except Exception as e2:
                raise ProcessorError(f"Text extraction failed: {str(e2)}")

    def _detect_document_type(
        self,
        text: str,
        metadata: Dict[str, Any]
    ) -> str:
        """
        Detect document type from content

        Returns:
            Document type: invoice, receipt, report, letter, etc.
        """
        text_lower = text.lower()

        # Invoice/Receipt detection
        invoice_keywords = ['invoice', 'bill to', 'total amount', 'payment due', 'invoice number']
        receipt_keywords = ['receipt', 'paid', 'transaction', 'subtotal']

        if any(keyword in text_lower for keyword in invoice_keywords):
            return 'invoice'
        elif any(keyword in text_lower for keyword in receipt_keywords):
            return 'receipt'
        elif 'contract' in text_lower:
            return 'contract'
        elif any(word in text_lower for word in ['report', 'analysis', 'summary']):
            return 'report'
        else:
            return 'document'

    async def _extract_invoice_data(self, text: str) -> Optional[Dict[str, Any]]:
        """
        Extract structured data from invoice/receipt using LLM

        Args:
            text: Document text

        Returns:
            Dictionary with invoice data
        """
        try:
            prompt = f"""Extract structured data from this invoice/receipt.

Invoice/Receipt Text:
{text[:2000]}

Return JSON with these fields (use null if not found):
{{
    "invoice_number": "...",
    "date": "...",
    "vendor": "...",
    "total_amount": "...",
    "currency": "...",
    "items": ["item1", "item2"]
}}"""

            response = await self.llm.generate(prompt, max_tokens=500, temperature=0.1)

            # Parse JSON
            import json
            if "```json" in response:
                response = response.split("```json")[1].split("```")[0].strip()
            elif "```" in response:
                response = response.split("```")[1].split("```")[0].strip()

            return json.loads(response)

        except Exception as e:
            print(f"Invoice data extraction failed: {e}")
            return None

    def _extract_key_points(
        self,
        text: str,
        metadata: Dict[str, Any]
    ) -> List[str]:
        """Extract key points from document"""
        # Extract first few paragraphs
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]

        # Get first sentence of each paragraph (up to 5)
        key_points = []
        for para in paragraphs[:5]:
            sentences = para.split('.')
            if sentences:
                key_points.append(sentences[0].strip())

        return key_points
